import requests
import json
import docx
import sys
from PyPDF2 import PdfFileReader

file = sys.argv[1]
tar = sys.argv[2]


IAM_TOKEN = 't1.9euelZqLmZWXjsuPxpWMmJyMz5HLie3rnpWakoqTnMeXkJWcm4nHk5mdlZbl8_cXBTZe-e8GLklj_N3z91czM1757wYuSWP8.TVRyb20UDZni-MFIDyVFnWGt6HsKr1y_b76rFC8v9Kl6CKqQCc4cQlQrqH0RIQc0RoodpbWjZrAP3YxNzSJ0DQ'
folder_id = 'b1gsueshknlj50pvtij1'
target_language = f'{tar}'

file_path = '.\source'
pdf_document = f'{file_path}\{file}.pdf'
print(pdf_document)
doc = docx.Document()

with open(pdf_document, "rb") as filehandle:
    pdf = PdfFileReader(filehandle)
    pages = pdf.getNumPages()
    for i in range(pages):
        page = pdf.getPage(i)
        doc.add_paragraph(page.extractText())


all_paras = doc.paragraphs
translated_doc = docx.Document()

for para in all_paras:
    text = para.text
    body = {
        "targetLanguageCode": target_language,
        "texts": text,
        "folderId": folder_id,
    }

    headers = {
        "Content-Type": "application/json",
        "Authorization": "Bearer {0}".format(IAM_TOKEN)
    }

    response = requests.post('https://translate.api.cloud.yandex.net/translate/v2/translate',
        json = body,
        headers = headers
    )

    json_load = json.loads(response.text)
    if list(json_load.keys())[0] != 'translations':
        pass
    else:
        translated_doc.add_paragraph(json_load['translations'][0]['text'])
translated_doc.save(f'{file_path}\{file}_{tar}_translated.docx')
